"""
ExportService — Moteur de rendu multi-format
Formats supportés : PDF, DOCX, XLSX, HTML, MD
"""
import io
import re
import logging
import tempfile
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _parse_markdown_tables(md: str) -> List[Dict[str, Any]]:
    """Extrait les tableaux Markdown sous forme de liste de dicts {headers, rows}."""
    tables = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        # Ligne header (contient |)
        if "|" in lines[i]:
            header_line = lines[i]
            # Ligne separator suivante
            if i + 1 < len(lines) and re.match(r"^\s*\|[-| :]+\|\s*$", lines[i + 1]):
                headers = [h.strip() for h in header_line.strip().strip("|").split("|")]
                rows = []
                j = i + 2
                while j < len(lines) and "|" in lines[j]:
                    row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                    rows.append(row)
                    j += 1
                tables.append({"headers": headers, "rows": rows, "start": i, "end": j})
                i = j
                continue
        i += 1
    return tables


def _md_to_html(md: str, title: str = "Export") -> str:
    """Convertit Markdown en HTML avec styles inline pour rendu autonome."""
    try:
        import markdown as md_lib
        body = md_lib.markdown(
            md,
            extensions=["tables", "fenced_code", "codehilite", "toc", "nl2br"],
        )
    except ImportError:
        # Fallback minimal si markdown non installé
        body = "<pre>" + md.replace("<", "&lt;").replace(">", "&gt;") + "</pre>"

    now = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{title}</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
      font-size: 15px;
      line-height: 1.7;
      color: #1a1a2e;
      background: #ffffff;
      max-width: 860px;
      margin: 0 auto;
      padding: 40px 32px 60px;
    }}
    h1, h2, h3, h4 {{
      font-weight: 700;
      margin-top: 1.6em;
      margin-bottom: 0.5em;
      color: #111827;
    }}
    h1 {{ font-size: 2em; border-bottom: 2px solid #e5e7eb; padding-bottom: 0.3em; }}
    h2 {{ font-size: 1.5em; border-bottom: 1px solid #e5e7eb; padding-bottom: 0.2em; }}
    h3 {{ font-size: 1.2em; }}
    p {{ margin: 0.8em 0; }}
    code {{
      background: #f3f4f6;
      border-radius: 4px;
      padding: 2px 6px;
      font-size: 0.88em;
      font-family: 'JetBrains Mono', 'Fira Code', monospace;
      color: #374151;
    }}
    pre {{
      background: #1e293b;
      color: #e2e8f0;
      border-radius: 8px;
      padding: 16px 20px;
      overflow-x: auto;
      font-size: 0.86em;
      line-height: 1.6;
    }}
    pre code {{ background: none; color: inherit; padding: 0; }}
    blockquote {{
      border-left: 4px solid #6366f1;
      margin: 1em 0;
      padding: 8px 16px;
      background: #f5f3ff;
      border-radius: 0 6px 6px 0;
      color: #4b5563;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 1.2em 0;
      font-size: 0.92em;
    }}
    th {{
      background: #6366f1;
      color: white;
      padding: 10px 14px;
      text-align: left;
      font-weight: 600;
    }}
    td {{
      padding: 9px 14px;
      border-bottom: 1px solid #e5e7eb;
    }}
    tr:nth-child(even) td {{ background: #f9fafb; }}
    tr:hover td {{ background: #f0f0ff; }}
    a {{ color: #6366f1; text-decoration: none; }}
    a:hover {{ text-decoration: underline; }}
    ul, ol {{ padding-left: 1.6em; }}
    li {{ margin: 0.3em 0; }}
    hr {{ border: none; border-top: 1px solid #e5e7eb; margin: 2em 0; }}
    .export-meta {{
      font-size: 0.78em;
      color: #9ca3af;
      border-top: 1px solid #e5e7eb;
      margin-top: 3em;
      padding-top: 1em;
    }}
  </style>
</head>
<body>
  {body}
  <div class="export-meta">Exported on {now}</div>
</body>
</html>"""


# ── ExportService ──────────────────────────────────────────────────────────────

class ExportService:
    """
    Service d'export multi-format à partir de contenu Markdown.

    Usage :
        svc = ExportService()
        buf, mime, ext = await svc.render(content="# Hello", fmt="pdf", title="My doc")
        # buf → BytesIO prêt pour StreamingResponse ou base64
    """

    SUPPORTED_FORMATS = ("pdf", "docx", "xlsx", "html", "md")

    async def render(
        self,
        content: str,
        fmt: str,
        title: str = "Export",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> tuple[io.BytesIO, str, str]:
        """
        Rendu du contenu dans le format demandé.

        Returns:
            (buffer, mime_type, file_extension)
        """
        fmt = fmt.lower().strip(".")
        if fmt not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {fmt}. Choose from {self.SUPPORTED_FORMATS}")

        dispatch = {
            "pdf":  self._render_pdf,
            "docx": self._render_docx,
            "xlsx": self._render_xlsx,
            "html": self._render_html,
            "md":   self._render_md,
        }
        return await dispatch[fmt](content, title, metadata or {})

    # ── PDF ────────────────────────────────────────────────────────────────────

    async def _render_pdf(
        self, content: str, title: str, metadata: Dict
    ) -> tuple[io.BytesIO, str, str]:
        try:
            from weasyprint import HTML as WP_HTML, CSS
        except ImportError as e:
            raise ImportError(
                "weasyprint is required for PDF export. "
                "Install it: pip install weasyprint"
            ) from e

        html_str = _md_to_html(content, title)
        buf = io.BytesIO()

        # WeasyPrint: HTML string → PDF bytes
        WP_HTML(string=html_str).write_pdf(buf)
        buf.seek(0)
        return buf, "application/pdf", "pdf"

    # ── DOCX ───────────────────────────────────────────────────────────────────

    async def _render_docx(
        self, content: str, title: str, metadata: Dict
    ) -> tuple[io.BytesIO, str, str]:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Styles de base
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Métadonnées document
        doc.core_properties.title = title
        doc.core_properties.created = datetime.utcnow()
        if metadata.get("author"):
            doc.core_properties.author = metadata["author"]

        # Parse Markdown ligne par ligne
        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            # Headings
            if line.startswith("### "):
                p = doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                p = doc.add_heading(line[2:].strip(), level=1)

            # Separator
            elif line.strip() in ("---", "***", "___"):
                doc.add_paragraph().add_run().add_break()

            # Code block
            elif line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                p = doc.add_paragraph("\n".join(code_lines))
                p.style = doc.styles["No Spacing"]
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

            # Blockquote
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:].strip())
                p.paragraph_format.left_indent = Inches(0.4)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
                    run.font.italic = True

            # Unordered list
            elif re.match(r"^[-*+] ", line):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")

            # Ordered list
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line).strip(), style="List Number")

            # Table (detect | row)
            elif "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[-| :]+\|\s*$", lines[i + 1]):
                headers = [h.strip() for h in line.strip().strip("|").split("|")]
                i += 2  # skip separator
                rows = []
                while i < len(lines) and "|" in lines[i]:
                    rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                    i += 1
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                # Header row
                hdr_cells = table.rows[0].cells
                for col_idx, h in enumerate(headers):
                    hdr_cells[col_idx].text = h
                    for paragraph in hdr_cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                # Data rows
                for row_idx, row in enumerate(rows):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, cell_val in enumerate(row):
                        if col_idx < len(row_cells):
                            row_cells[col_idx].text = cell_val
                continue  # i already incremented

            # Blank line
            elif not line.strip():
                pass

            # Regular paragraph with inline bold/italic
            else:
                p = doc.add_paragraph()
                # Parse inline **bold** and *italic*
                remaining = line
                while remaining:
                    bold_match = re.search(r"\*\*(.+?)\*\*", remaining)
                    italic_match = re.search(r"\*(.+?)\*", remaining)
                    if bold_match and (not italic_match or bold_match.start() <= italic_match.start()):
                        p.add_run(remaining[: bold_match.start()])
                        p.add_run(bold_match.group(1)).bold = True
                        remaining = remaining[bold_match.end():]
                    elif italic_match:
                        p.add_run(remaining[: italic_match.start()])
                        p.add_run(italic_match.group(1)).italic = True
                        remaining = remaining[italic_match.end():]
                    else:
                        p.add_run(remaining)
                        break

            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

    # ── XLSX ───────────────────────────────────────────────────────────────────

    async def _render_xlsx(
        self, content: str, metadata: Dict, title: str
    ) -> tuple[io.BytesIO, str, str]:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = openpyxl.Workbook()
        wb.properties.title = title

        tables = _parse_markdown_tables(content)

        if tables:
            # One sheet per table
            for t_idx, table in enumerate(tables):
                ws_title = f"Table {t_idx + 1}" if t_idx > 0 else "Data"
                if t_idx == 0:
                    ws = wb.active
                    ws.title = ws_title
                else:
                    ws = wb.create_sheet(ws_title)

                header_fill = PatternFill("solid", fgColor="6366F1")
                header_font = Font(bold=True, color="FFFFFF", size=11)
                border_side = Side(style="thin", color="E5E7EB")
                cell_border = Border(
                    left=border_side, right=border_side,
                    top=border_side, bottom=border_side
                )

                # Headers
                for col_idx, header in enumerate(table["headers"], start=1):
                    cell = ws.cell(row=1, column=col_idx, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border = cell_border

                # Rows
                for row_idx, row in enumerate(table["rows"], start=2):
                    alt_fill = PatternFill("solid", fgColor="F9FAFB") if row_idx % 2 == 0 else None
                    for col_idx, value in enumerate(row, start=1):
                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                        cell.border = cell_border
                        if alt_fill:
                            cell.fill = alt_fill

                # Auto column width
                for col in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col), default=0)
                    ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

                ws.freeze_panes = "A2"
        else:
            # No table found → dump raw content into single column
            ws = wb.active
            ws.title = "Content"
            ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=13)
            for line_idx, line in enumerate(content.splitlines(), start=2):
                ws.cell(row=line_idx, column=1, value=line)
            ws.column_dimensions["A"].width = 100

        # Summary sheet
        meta_ws = wb.create_sheet("_meta")
        meta_ws.cell(1, 1, "Generated").font = Font(bold=True)
        meta_ws.cell(1, 2, datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))
        meta_ws.cell(2, 1, "Title").font = Font(bold=True)
        meta_ws.cell(2, 2, title)
        meta_ws.column_dimensions["A"].width = 16
        meta_ws.column_dimensions["B"].width = 40

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"

    # ── HTML ───────────────────────────────────────────────────────────────────

    async def _render_html(
        self, content: str, title: str, metadata: Dict
    ) -> tuple[io.BytesIO, str, str]:
        html_str = _md_to_html(content, title)
        buf = io.BytesIO(html_str.encode("utf-8"))
        return buf, "text/html; charset=utf-8", "html"

    # ── MD ─────────────────────────────────────────────────────────────────────

    async def _render_md(
        self, content: str, title: str, metadata: Dict
    ) -> tuple[io.BytesIO, str, str]:
        header = f"# {title}\n\n_Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}_\n\n---\n\n"
        buf = io.BytesIO((header + content).encode("utf-8"))
        return buf, "text/markdown; charset=utf-8", "md"